# from flask import Flask, request, jsonify, send_from_directory
# import os
# import vertexai
# from vertexai.generative_models import GenerativeModel, ChatSession, Content, Part
# from google.cloud import storage
# import logging
# from docx import Document

# app = Flask(__name__, static_folder='.', static_url_path='')

# # Set up logging
# logging.basicConfig(level=logging.INFO)
# logger = logging.getLogger(__name__)

# # Initialize Vertex AI
# project_id = "black-cirrus-461305-f6"
# location = "us-east4"
# vertexai.init(project=project_id, location=location)

# # Initialize Cloud Storage client
# storage_client = storage.Client(project=project_id)

# # Load Gemini model
# try:
#     model = GenerativeModel("gemini-1.5-pro")
#     logger.info("Successfully loaded Gemini model: gemini-1.5-pro")
# except Exception as e:
#     logger.error(f"Failed to load Gemini model: {str(e)}")
#     raise

# def load_docx_from_storage(bucket_name, file_name):
#     try:
#         bucket = storage_client.bucket(bucket_name)
#         blob = bucket.blob(file_name)
#         docx_bytes = blob.download_as_bytes()

#         with open("/tmp/temp.docx", "wb") as f:
#             f.write(docx_bytes)

#         document = Document("/tmp/temp.docx")
#         text = "\n".join([para.text for para in document.paragraphs if para.text.strip()])
#         logger.info(f"Successfully loaded DOCX from {file_name}")
#         return text
#     except Exception as e:
#         logger.error(f"Failed to load DOCX from {bucket_name}/{file_name}: {str(e)}")
#         return None


# @app.route('/')
# def serve_index():
#     return send_from_directory('.', 'index.html')

# @app.route('/chat', methods=['POST'])
# def chat():
#     data = request.get_json()
#     if not data or 'message' not in data:
#         logger.error("Invalid request: No message provided")
#         return jsonify({"error": "Please provide a message"}), 400

#     user_message = data['message']
#     tenant_id = request.headers.get("x-tenant-id") or data.get("tenant_id")
#     if not tenant_id:
#         return jsonify({"error": "Missing tenant_id"}), 400

#     logger.info(f"Received message from tenant '{tenant_id}': {user_message}")

#     # Load DOCX file for the tenant
#     BUCKET_NAME = "multi-tenant-ex"
#     DOCX_FILE_NAME = f"{tenant_id}/digital_wall_chart.docx"
#     doc_text = load_docx_from_storage(BUCKET_NAME, DOCX_FILE_NAME)

#     if not doc_text:
#         return jsonify({"response": f"Error: No document found for tenant '{tenant_id}'"}), 500

#     # Build prompt
#     prompt = f"User query: {user_message}\n\nDocument Content:\n{doc_text}\n\nAnswer the query based on the document content."

#     # Start Gemini chat session
#     try:
#         chat_session = model.start_chat(
#             history=[
#                 Content(role="user", parts=[Part.from_text("What is the capital of France?")]),
#                 Content(role="model", parts=[Part.from_text("The capital of France is Paris.")])
#             ]
#         )
#         logger.info("Gemini chat session started")
#     except Exception as e:
#         logger.error(f"Failed to start chat: {str(e)}")
#         return jsonify({"response": f"Error starting chat: {str(e)}"}), 500


#     # Send message to Gemini
#     try:
#         response = chat_session.send_message(
#             content=prompt,
#             generation_config={
#                 "temperature": 0.2,
#                 "max_output_tokens": 256,
#                 "top_p": 0.8,
#                 "top_k": 40
#             }
#         )
#         bot_response = response.text
#         logger.info(f"Gemini response: {bot_response}")
#     except Exception as e:
#         logger.error(f"Failed to get Gemini response: {str(e)}")
#         bot_response = f"Error: {str(e)}"

#     return jsonify({"response": bot_response})


# if __name__ == '__main__':
#     port = int(os.environ.get('PORT', 8080))
#     app.run(host='0.0.0.0', port=port)

from flask import Flask, request, jsonify, send_from_directory
import os
import vertexai
from vertexai.generative_models import GenerativeModel, ChatSession, Content, Part
from google.cloud import storage
import logging
from docx import Document

app = Flask(__name__, static_folder='.', static_url_path='')

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Configuration
CONFIG = {
    "project_id": "black-cirrus-461305-f6",
    "location": "us-east4",
    "bucket_name": "multi-tenant-ex",
    "tenants": {
        "tenant1": {
            "datastore_id": "tenant1_1750145181743",
            "docx_file": "tenant1/digital_wall_chart.docx",
            "app_id": "multitenantfaqbot_1750145664170"
        },
        "tenant2": {
            "datastore_id": "tenant2_1750145004762",
            "docx_file": "tenant2/digital_wall_chart.docx",
            "app_id": "multitenantfaqbot_1750145664170"
        }
    }
}

# Initialize Vertex AI - WITH ERROR HANDLING
try:
    vertexai.init(project=CONFIG["project_id"], location=CONFIG["location"])
    logger.info(f"Vertex AI initialized in {CONFIG['location']}")
except Exception as e:
    logger.error(f"Failed to initialize Vertex AI: {str(e)}")
    raise

# Initialize Cloud Storage client
storage_client = storage.Client(project=CONFIG["project_id"])

# Load Gemini model - UPDATED MODEL NAME
try:
    model = GenerativeModel("gemini-1.0-pro")  # Verified working model
    logger.info("Successfully loaded Gemini model: gemini-1.0-pro")
except Exception as e:
    logger.error(f"Failed to load Gemini model: {str(e)}")
    raise

# Initialize Vertex AI
vertexai.init(project=CONFIG["project_id"], location=CONFIG["location"])

# Initialize Cloud Storage client
storage_client = storage.Client(project=CONFIG["project_id"])

# Load Gemini model
try:
    model = GenerativeModel("gemini-1.0-pro")
    logger.info("Successfully loaded Gemini model: gemini-1.0-pro")
except Exception as e:
    logger.error(f"Failed to load Gemini model: {str(e)}")
    raise

def get_tenant_config(tenant_id):
    """Get configuration for a specific tenant"""
    if tenant_id not in CONFIG["tenants"]:
        logger.error(f"Unknown tenant ID: {tenant_id}")
        return None
    return CONFIG["tenants"][tenant_id]

def load_docx_from_storage(bucket_name, file_name):
    try:
        bucket = storage_client.bucket(bucket_name)
        blob = bucket.blob(file_name)
        docx_bytes = blob.download_as_bytes()

        with open("/tmp/temp.docx", "wb") as f:
            f.write(docx_bytes)

        document = Document("/tmp/temp.docx")
        text = "\n".join([para.text for para in document.paragraphs if para.text.strip()])
        logger.info(f"Successfully loaded DOCX from {file_name}")
        return text
    except Exception as e:
        logger.error(f"Failed to load DOCX from {bucket_name}/{file_name}: {str(e)}")
        return None

@app.route('/')
def serve_index():
    return send_from_directory('.', 'index.html')

@app.route('/chat', methods=['POST'])
def chat():
    data = request.get_json()
    if not data or 'message' not in data:
        logger.error("Invalid request: No message provided")
        return jsonify({"error": "Please provide a message"}), 400

    user_message = data['message']
    tenant_id = request.headers.get("x-tenant-id") or data.get("tenant_id")
    
    if not tenant_id:
        return jsonify({"error": "Missing tenant_id"}), 400

    logger.info(f"Received message from tenant '{tenant_id}': {user_message}")

    # Get tenant-specific configuration
    tenant_config = get_tenant_config(tenant_id)
    if not tenant_config:
        return jsonify({"error": f"Invalid tenant ID: {tenant_id}"}), 400

    # Load tenant-specific DOCX file
    doc_text = load_docx_from_storage(CONFIG["bucket_name"], tenant_config["docx_file"])

    if not doc_text:
        return jsonify({"response": f"Error: No document found for tenant '{tenant_id}'"}), 500

    # Build tenant-specific prompt
    prompt = f"""You are a chatbot for tenant {tenant_id}. 
    User query: {user_message}
    
    Tenant-specific document content:
    {doc_text}
    
    Answer the query based on the document content, specifically for this tenant."""

    # Start Gemini chat session with tenant context
    try:
        chat_session = model.start_chat(
            history=[
                Content(role="user", parts=[Part.from_text(f"You are a chatbot for {tenant_id}")]),
                Content(role="model", parts=[Part.from_text(f"Understood. I will respond as a chatbot specifically for {tenant_id}")])
            ]
        )
        logger.info(f"Gemini chat session started for tenant {tenant_id}")
    except Exception as e:
        logger.error(f"Failed to start chat: {str(e)}")
        return jsonify({"response": f"Error starting chat: {str(e)}"}), 500

    # Send message to Gemini
    try:
        response = chat_session.send_message(
            content=prompt,
            generation_config={
                "temperature": 0.2,
                "max_output_tokens": 256,
                "top_p": 0.8,
                "top_k": 40
            }
        )
        bot_response = response.text
        logger.info(f"Gemini response for tenant {tenant_id}: {bot_response}")
    except Exception as e:
        logger.error(f"Failed to get Gemini response: {str(e)}")
        bot_response = f"Error: {str(e)}"

    return jsonify({
        "response": bot_response,
        "tenant_id": tenant_id,
        "source": tenant_config["docx_file"]
    })

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 8080))
    app.run(host='0.0.0.0', port=port)
